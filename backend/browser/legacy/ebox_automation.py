"""
E-Box Course Automation Tool - REVISED FOR PROJECT-BASED LEARNING
Automates course completion on https://pro.e-box.co.in

ACTUAL E-BOX STRUCTURE (based on screenshots):
- Course Page: Shows Topics (left sidebar) and Sections (tabs: i-Learn, i-Explore, i-Analyse, i-Design)
- Topics = Unit names (e.g., "Solution Of Ordinary Differential...", "Vector Calculus")
- Sections = Tabs (i-Learn, i-Explore, i-Analyse, i-Design)
- Each Section has a Project with problems to solve
- Projects require code solutions with command line arguments
"""

import logging
import asyncio
import re
from typing import Annotated, Dict, List, Any
from livekit.agents import RunContext, function_tool

logger = logging.getLogger(__name__)

def _extract_json_from_text(text: str) -> dict:
    import json
    text = text.strip()
    match = re.search(r'```(?:json)?\s*(.*?)\s*```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
            
    match = re.search(r'(\{.*\})', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
            
    try:
        # Sometimes there's unescaped characters, let's try direct parse
        return json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"[JSON] Failed to parse JSON: {e} \nRaw Text: {text[:200]}...")
        raise

# Hardcoded credentials
EBOX_USERNAME = "SIT25CS170"
EBOX_PASSWORD = "SIT25CS170"
EBOX_LOGIN_URL = "https://pro.e-box.co.in/login"

# Course name mappings
COURSE_MAPPINGS = {
    # Differential Equations
    "differential": "Differential Equations And Complex Analysis",
    "diff eq": "Differential Equations And Complex Analysis",
    "differential equations": "Differential Equations And Complex Analysis",
    "complex analysis": "Differential Equations And Complex Analysis",
    "maths": "Differential Equations And Complex Analysis",
    "math": "Differential Equations And Complex Analysis",
    "mathematics": "Differential Equations And Complex Analysis",
    # Biology
    "biology": "Biology for Engineers",
    "bio": "Biology for Engineers",
    "biology for engineers": "Biology for Engineers",
    # Chemistry
    "chemistry": "Principles of Chemistry in Engineering",
    "chem": "Principles of Chemistry in Engineering",
    "principles of chemistry": "Principles of Chemistry in Engineering",
    # Digital Electronics / Logic
    "digital electronics": "Digital Electronics",
    "digital logic": "Digital Logic and Design",
    "digital": "Digital Electronics",
    "logic design": "Digital Logic and Design",
    "de": "Digital Electronics",
    # Python / Programming
    "python": "Programming Fundamentals using Python",
    "programming": "Programming Fundamentals using Python",
    "programming fundamentals": "Programming Fundamentals using Python",
    "pfp": "Programming Fundamentals using Python",
    # Physics / Engineering Physics
    "physics": "Engineering Physics",
    "engineering physics": "Engineering Physics",
    "ep": "Engineering Physics",
    # Electrical
    "electrical": "Basic Electrical Engineering",
    "bee": "Basic Electrical Engineering",
    "basic electrical": "Basic Electrical Engineering",
    # Electronics
    "electronics": "Basic Electronics Engineering",
    "beee": "Basic Electronics Engineering",
    "basic electronics": "Basic Electronics Engineering",
    # Communication
    "communication": "Technical Communication",
    "english": "Technical Communication",
    "tc": "Technical Communication",
    # Problem solving
    "problem solving": "Problem Solving Using C",
    "c programming": "Problem Solving Using C",
    "c language": "Problem Solving Using C",
    "psc": "Problem Solving Using C",
    # Environmental Science
    "environmental": "Environmental Sciences",
    "evs": "Environmental Sciences",
    "environment": "Environmental Sciences",
}

# Section/Phase names (in order they appear on the site)
SECTIONS = ["i-Learn", "i-Explore", "i-Analyse", "i-Design", "i-Assess"]

# Real topic lists per course (discovered from live e-box session)
COURSE_TOPICS: dict[str, list[str]] = {
    "Differential Equations And Complex Analysis": [
        "Solution Of Ordinary Differential Equations",
        "Vector Calculus",
        "Analytic Functions",
        "Complex Integration",
        "Partial Differential Equation",
    ],
    "Biology for Engineers": [
        "Introduction to Biology",
        "Cell Biology",
        "Genetics",
        "Ecology",
        "Biotechnology Applications",
    ],
    "Principles of Chemistry in Engineering": [
        "Atomic Structure and Bonding",
        "Thermodynamics",
        "Electrochemistry",
        "Polymers",
        "Corrosion",
    ],
    "Digital Electronics": [
        "Fundamentals of Digital Logic and Boolean Algebra",
        "Boolean Function Simplification and Circuit Design",
        "Design of Combinational Circuits",
        "Sequential Circuits and Storage Elements",
        "Asynchronous Sequential Logic",
        "Memory Device And Digital Integrated Circuits",
    ],
    "Digital Logic and Design": [
        "Number Systems and Codes",
        "Boolean Algebra",
        "Combinational Logic Design",
        "Sequential Logic Design",
        "Memory and Programmable Logic",
    ],
    "Programming Fundamentals using Python": [
        "Introduction to Python",
        "Control Structures",
        "Functions and Modules",
        "Data Structures",
        "File Handling and OOP",
    ],
    "Problem Solving Using C": [
        "Introduction to Programming",
        "Control Flow",
        "Arrays and Strings",
        "Functions and Pointers",
        "Structures and File I/O",
    ],
    "Technical Communication": [
        "Introduction to Professional Communication",
        "Parliamentary English and Formal Usage",
        "Pronunciation, Stress, and Intonation",
        "Words Often Confused & Common Errors",
        "Principles of Communicative English",
        "Lucid Writing",
        "Individual and Team Communication Skills",
        "Grammar for Professionals – Relative Clauses",
        "Life Skills and Interpersonal Skills",
        "Professional Ethics and Code of Conduct",
    ],
}


def extract_course_and_topic(user_input: str) -> Dict[str, Any]:
    """
    Extract course name and topic from natural language input.
    """
    input_lower = user_input.lower().strip()

    result = {
        "course": None,
        "topic": None,
        "section": None,  # i-Learn, i-Explore, i-Analyse, i-Design
        "all": False,
    }

    # Check if requesting all courses/topics
    if any(phrase in input_lower for phrase in ["finish course", "complete course", "do course", "finish my course", "complete all"]):
        result["all"] = True

    # Extract section/phase
    section_patterns = {
        "learn": "i-Learn",
        "explore": "i-Explore",
        "analyse": "i-Analyse",
        "analyze": "i-Analyse",
        "design": "i-Design",
    }
    for key, value in section_patterns.items():
        if key in input_lower:
            result["section"] = value
            break

    # Extract course name
    for key, value in COURSE_MAPPINGS.items():
        if key in input_lower:
            result["course"] = value
            break

    # Extract specific topic (if mentioned)
    topic_keywords = [
        "solution of ordinary differential equations",
        "vector calculus",
        "analytic functions",
        "complex integration",
        "partial differential equation"
    ]
    for topic in topic_keywords:
        if topic in input_lower:
            result["topic"] = topic.title()
            break

    return result


async def login_to_ebox(browser) -> Dict[str, Any]:
    """
    Login to E-Box using hardcoded credentials.
    E-Box always requires login - this auto-fills the form every time.
    """
    try:
        logger.info("[EBox] Navigating to login page...")

        # Navigate to login page
        nav_result = await browser.navigate(EBOX_LOGIN_URL)
        if not nav_result.get("success"):
            return {"success": False, "error": f"Failed to load page: {nav_result.get('error')}"}

        await asyncio.sleep(2)

        # Check if already on dashboard (rare but possible)
        current_url = await browser.get_current_url()
        if "login" not in current_url.lower():
            logger.info("[EBox] Already logged in!")
            return {"success": True, "message": "Already logged in"}

        logger.info(f"[EBox] Filling login form with username: {EBOX_USERNAME}")

        page = browser.page

        try:
            await page.wait_for_load_state("networkidle", timeout=10000)

            # Fill username
            username_selectors = ['input[name="username"]', 'input[type="text"]', '#username']
            for selector in username_selectors:
                try:
                    element = await page.query_selector(selector)
                    if element:
                        await element.fill(EBOX_USERNAME)
                        logger.info(f"[EBox] Filled username using: {selector}")
                        break
                except:
                    continue

            # Fill password
            password_selectors = ['input[name="password"]', 'input[type="password"]', '#password']
            for selector in password_selectors:
                try:
                    element = await page.query_selector(selector)
                    if element:
                        await element.fill(EBOX_PASSWORD)
                        logger.info(f"[EBox] Filled password using: {selector}")
                        break
                except:
                    continue

            await asyncio.sleep(0.5)

            # Click submit
            submit_selectors = ['button[type="submit"]', 'input[type="submit"]', 'button:has-text("Login")']
            for selector in submit_selectors:
                try:
                    element = await page.query_selector(selector)
                    if element:
                        await element.click()
                        logger.info(f"[EBox] Clicked submit using: {selector}")
                        break
                except:
                    continue

            await asyncio.sleep(3)

            current_url = await browser.get_current_url()
            if "login" not in current_url.lower():
                logger.info("[EBox] Login successful!")
                return {"success": True, "message": "Logged in successfully!"}
            else:
                return {"success": False, "error": "Login failed - still on login page"}

        except Exception as form_error:
            logger.error(f"[EBox] Form fill error: {form_error}")
            return {"success": False, "error": f"Could not fill login form: {form_error}"}

    except Exception as e:
        logger.error(f"[EBox] Login error: {e}")
        return {"success": False, "error": str(e)}


async def navigate_to_course(browser, course_name: str) -> Dict[str, Any]:
    """Navigate to a specific course from the dashboard."""
    try:
        logger.info(f"[EBox] Navigating to course: {course_name}")
        page = browser.page
        await asyncio.sleep(2)

        # Build search terms: full name, first word, each individual word
        words = course_name.split()
        search_terms = [course_name] + words[:2]

        # Try broad set of selectors that could hold course cards
        selectors = [
            'a[href*="course"]',
            'a[href*="Course"]',
            '[class*="course"] a',
            '[class*="card"] a',
            '.item a',
            'a',
        ]

        for selector in selectors:
            try:
                links = await page.query_selector_all(selector)
            except Exception:
                continue
            logger.info(f"[EBox] Trying selector '{selector}' — {len(links)} elements found")
            for term in search_terms:
                for link in links:
                    try:
                        text = (await link.text_content() or "").strip()
                        if text and term.lower() in text.lower():
                            logger.info(f"[EBox] Clicking course link (term='{term}'): {text[:80]}")
                            await link.click()
                            await asyncio.sleep(3)
                            current = await browser.get_current_url()
                            if "login" not in current.lower():
                                return {"success": True, "course": course_name, "url": current}
                    except Exception:
                        continue

        # Last resort: use Playwright locator by visible text
        for term in search_terms:
            try:
                loc = page.get_by_text(re.compile(term, re.IGNORECASE)).first
                await loc.click(timeout=3000)
                await asyncio.sleep(3)
                logger.info(f"[EBox] Clicked course via text locator: {term}")
                return {"success": True, "course": course_name}
            except Exception:
                continue

        return {"success": False, "error": f"Course '{course_name}' not found on dashboard"}

    except Exception as e:
        logger.error(f"[EBox] Course navigation error: {e}")
        return {"success": False, "error": str(e)}


async def get_topics_list(browser) -> Dict[str, Any]:
    """Get list of topics from the left sidebar."""
    try:
        page = browser.page
        await asyncio.sleep(1)

        # Topics are in the left sidebar - look for the Topics section
        topics = []

        # Try to find topic elements in sidebar
        topic_selectors = [
            '.topic-item',
            '[class*="topic"]',
            '.sidebar a',
            'a:has-text("Solution")',
            'a:has-text("Vector")',
            'a:has-text("Analytic")',
            'a:has-text("Complex")',
            'a:has-text("Partial")',
        ]

        # Get all text that looks like topic names
        body_text = await page.inner_text('body')

        # All known topics across all courses
        all_known_topics = [
            topic
            for topic_list in COURSE_TOPICS.values()
            for topic in topic_list
        ] + [
            # Extra fallback partial matches
            "Solution Of Ordinary Differential",
            "Vector Calculus",
            "Analytic Functions",
            "Complex Integration",
            "Partial Differential Equation",
            "Fundamentals of Digital Logic and Boolean Algebra",
            "Boolean Function Simplification and Circuit Design",
            "Design of Combinational Circuits",
            "Sequential Circuits and Storage Elements",
            "Asynchronous Sequential Logic",
            "Memory Device And Digital Integrated Circuits",
        ]

        seen = set()
        for topic in all_known_topics:
            if topic.lower() in body_text.lower() and topic not in seen:
                topics.append(topic)
                seen.add(topic)

        logger.info(f"[EBox] Found topics: {topics}")
        return {"success": True, "topics": topics}

    except Exception as e:
        logger.error(f"[EBox] Error getting topics: {e}")
        return {"success": False, "error": str(e)}


async def click_topic(browser, topic_name: str) -> Dict[str, Any]:
    """Click on a topic in the left sidebar."""
    try:
        page = browser.page
        logger.info(f"[EBox] Clicking topic: {topic_name}")

        # Try to find and click the topic
        search_terms = [topic_name, topic_name.split()[0], topic_name[:20]]

        for term in search_terms:
            try:
                # Try clicking by text
                locator = page.get_by_text(re.compile(term, re.IGNORECASE)).first
                await locator.click()
                logger.info(f"[EBox] Clicked topic: {term}")
                await asyncio.sleep(2)
                return {"success": True, "topic": topic_name}
            except:
                continue

        return {"success": False, "error": f"Could not find topic: {topic_name}"}

    except Exception as e:
        logger.error(f"[EBox] Topic click error: {e}")
        return {"success": False, "error": str(e)}


async def click_section(browser, section_name: str) -> Dict[str, Any]:
    """Click on a section tab (i-Learn, i-Explore, i-Analyse, i-Design, i-Assess)."""
    try:
        page = browser.page
        logger.info(f"[EBox] Clicking section: {section_name}")

        # Based on exploration: sections are <a> tags with class "item" in a ".ui.pointing.secondary.menu"
        # Selector: a.item with text matching section name

        # Method 1: Click using the specific menu item selector
        try:
            section_link = await page.query_selector(f'a.item:has-text("{section_name}")')
            if section_link:
                await section_link.click()
                logger.info(f"[EBox] Clicked section using a.item selector: {section_name}")
                await asyncio.sleep(2)
                return {"success": True, "section": section_name}
        except Exception as e:
            logger.warning(f"[EBox] Method 1 failed: {e}")

        # Method 2: Find the menu and click the link inside it
        try:
            menu = await page.query_selector('.ui.pointing.secondary.menu')
            if menu:
                links = await menu.query_selector_all('a.item')
                for link in links:
                    text = await link.inner_text()
                    if section_name.lower() in text.lower():
                        await link.click()
                        logger.info(f"[EBox] Clicked section from menu: {section_name}")
                        await asyncio.sleep(2)
                        return {"success": True, "section": section_name}
        except Exception as e:
            logger.warning(f"[EBox] Method 2 failed: {e}")

        # Method 3: Direct text match
        try:
            await page.get_by_text(section_name, exact=True).first.click()
            logger.info(f"[EBox] Clicked section by text: {section_name}")
            await asyncio.sleep(2)
            return {"success": True, "section": section_name}
        except Exception as e:
            logger.warning(f"[EBox] Method 3 failed: {e}")

        # Method 4: Locator with text
        try:
            await page.locator(f'a:has-text("{section_name}")').first.click()
            logger.info(f"[EBox] Clicked section using locator: {section_name}")
            await asyncio.sleep(2)
            return {"success": True, "section": section_name}
        except Exception as e:
            logger.warning(f"[EBox] Method 4 failed: {e}")

        return {"success": False, "error": f"Could not find section: {section_name}"}

    except Exception as e:
        logger.error(f"[EBox] Section click error: {e}")
        return {"success": False, "error": str(e)}


async def click_project(browser) -> Dict[str, Any]:
    """Click on the project link to enter the problem page."""
    try:
        page = browser.page
        logger.info("[EBox] Looking for project to click...")

        # Wait a moment for content to load after section click
        await asyncio.sleep(1)

        # Based on exploration: Project links are <a> tags with href containing "attempt"
        # Example: <a href="/attempt/233693">Solution Of Ordinary Differential Equations / iLearn</a>

        # Method 1: Click link with href containing "attempt" (MOST RELIABLE)
        try:
            attempt_link = await page.query_selector('a[href*="attempt"]')
            if attempt_link:
                href = await attempt_link.get_attribute('href')
                text = await attempt_link.inner_text()
                logger.info(f"[EBox] Found attempt link: {text[:50]}... -> {href}")
                await attempt_link.click()
                await asyncio.sleep(3)
                return {"success": True, "href": href}
        except Exception as e:
            logger.warning(f"[EBox] Method 1 (attempt link) failed: {e}")

        # Method 2: Click link with href containing "project"
        try:
            project_link = await page.query_selector('a[href*="project"]')
            if project_link:
                await project_link.click()
                logger.info("[EBox] Clicked project link")
                await asyncio.sleep(3)
                return {"success": True}
        except Exception as e:
            logger.warning(f"[EBox] Method 2 failed: {e}")

        # Method 3: Click element containing "Completed" percentage
        try:
            # Look for the session content div and click it
            session_content = await page.query_selector('[class*="session_list_tab_content"]')
            if session_content:
                # Find the link inside
                link = await session_content.query_selector('a')
                if link:
                    await link.click()
                    logger.info("[EBox] Clicked link inside session content")
                    await asyncio.sleep(3)
                    return {"success": True}
        except Exception as e:
            logger.warning(f"[EBox] Method 3 failed: {e}")

        # Method 4: Click by text containing topic name and section
        try:
            await page.get_by_text(re.compile(r".*/(i-Learn|iLearn|i-Design|iDesign|i-Explore|i-Analyse)", re.IGNORECASE)).first.click()
            logger.info("[EBox] Clicked project by text pattern")
            await asyncio.sleep(3)
            return {"success": True}
        except Exception as e:
            logger.warning(f"[EBox] Method 4 failed: {e}")

        return {"success": False, "error": "Could not find project to click"}

    except Exception as e:
        logger.error(f"[EBox] Project click error: {e}")
        return {"success": False, "error": str(e)}
        logger.error(f"[EBox] Project click error: {e}")
        return {"success": False, "error": str(e)}


async def analyze_problem_page(browser) -> Dict[str, Any]:
    """Analyze the current problem page and extract problem details from DOM."""
    try:
        page = browser.page
        logger.info("[EBox] Analyzing problem page...")

        await asyncio.sleep(2)

        # Get page content - this is more reliable than vision
        body_text = await page.inner_text('body')

        # Extract problem details from the text
        problem_data = {
            "problem_title": "",
            "description": "",
            "context": "",
            "task": "",
            "input_format": "",
            "output_format": "",
            "sample_input": "",
            "sample_output": "",
            "programming_language": "Python"  # Default to Python
        }

        # Parse the body text to extract sections
        lines = body_text.split('\n')
        current_section = ""
        section_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect section headers
            if line in ["Context", "Problem Description", "Input Format", "Output Format",
                       "Sample Input", "Sample Output", "Model Definition", "Your task is to:",
                       "Learning Outcomes", "Sample Output Explanation"]:
                # Save previous section
                if current_section and section_content:
                    content = ' '.join(section_content)
                    if "context" in current_section.lower():
                        problem_data["context"] = content
                    elif "problem description" in current_section.lower():
                        problem_data["description"] = content
                    elif "input format" in current_section.lower():
                        problem_data["input_format"] = content
                    elif "output format" in current_section.lower():
                        problem_data["output_format"] = content
                    elif "sample input" in current_section.lower():
                        problem_data["sample_input"] = content
                    elif "sample output" in current_section.lower() and "explanation" not in current_section.lower():
                        problem_data["sample_output"] = content

                current_section = line
                section_content = []
            else:
                section_content.append(line)

        # Try to find the problem title (usually at the top)
        for line in lines[:20]:
            line = line.strip()
            if len(line) > 20 and len(line) < 150 and line[0].isupper():
                # Skip common non-title lines
                if line not in ["Context", "Problem Description", "Review", "Submit", "Next"]:
                    if not any(x in line.lower() for x in ["are widely", "one common", "you are given", "your task"]):
                        problem_data["problem_title"] = line
                        break

        # Build task description from what we found
        task_parts = []
        if problem_data["description"]:
            task_parts.append(problem_data["description"])
        if problem_data["context"]:
            task_parts.append(f"Context: {problem_data['context']}")

        problem_data["task"] = ' '.join(task_parts) if task_parts else body_text[:500]

        # Detect programming language from content
        if "python" in body_text.lower() or "matplotlib" in body_text.lower() or "numpy" in body_text.lower():
            problem_data["programming_language"] = "Python"
        elif "java" in body_text.lower():
            problem_data["programming_language"] = "Java"
        elif "#include" in body_text or "printf" in body_text:
            problem_data["programming_language"] = "C"

        logger.info(f"[EBox] Extracted problem: {problem_data['problem_title'][:50]}...")
        logger.info(f"[EBox] Language: {problem_data['programming_language']}")

        return {"success": True, "problem": problem_data}

    except Exception as e:
        logger.error(f"[EBox] Problem analysis error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


async def generate_solution_code(problem_data: Dict[str, Any]) -> Dict[str, Any]:
    """Generate solution code using AI."""
    try:
        import boto3
        import json

        bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')

        prompt = f"""
You are solving an E-Box programming problem that requires command line arguments.

PROBLEM: {problem_data.get('problem_title', 'Unknown')}

DESCRIPTION: {problem_data.get('description', '')}

TASK: {problem_data.get('task', '')}

INPUT FORMAT: {problem_data.get('input_format', 'Command line arguments')}

OUTPUT FORMAT: {problem_data.get('output_format', 'Print to stdout')}

SAMPLE INPUT: {problem_data.get('sample_input', '')}

SAMPLE OUTPUT: {problem_data.get('sample_output', '')}

LANGUAGE: {problem_data.get('programming_language', 'Python')}

CRITICAL E-BOX REQUIREMENTS:
1. The code MUST accept command line arguments using sys.argv (for Python) or equivalent
2. Parse the input format to determine what arguments are expected
3. E-Box provides a "Command line argument" textbox where test arguments are entered
4. E-Box will execute your code with these arguments: python script.py arg1 arg2 arg3
5. Generate realistic test arguments based on the sample input (these will be entered in the textbox)
6. Include import statements for required libraries (matplotlib, numpy, scipy, etc.)
7. Handle plotting/visualization if required by the problem

Generate complete, working code that E-Box can execute with command line arguments.

Return JSON:
{{
    "code": "complete working code with proper command line argument handling",
    "command_args": "space-separated arguments to test with (e.g. '100 0.4 10')",
    "explanation": "brief explanation of the solution approach"
}}

Return ONLY valid JSON.
"""

        response = bedrock.invoke_model(
            modelId='amazon.nova-pro-v1:0',
            body=json.dumps({
                "messages": [{"role": "user", "content": [{"text": prompt}]}],
                "inferenceConfig": {
                    "temperature": 0.3,
                    "max_new_tokens": 2000
                }
            })
        )

        response_body = json.loads(response['body'].read())
        result_text = response_body['output']['message']['content'][0]['text']

        result = _extract_json_from_text(result_text)

        return {
            "success": True,
            "code": result.get("code", ""),
            "command_args": result.get("command_args", ""),
            "explanation": result.get("explanation", "")
        }

    except Exception as e:
        logger.error(f"[EBox] Code generation error: {e}")
        return {"success": False, "error": str(e)}


async def submit_code(browser, code: str, command_args: str = "") -> Dict[str, Any]:
    """Submit code to E-Box code editor (ACE/Monaco editor)."""
    try:
        page = browser.page
        logger.info("[EBox] Submitting code...")

        # E-Box uses ACE editor - we need to click on the editor area first
        # The editor is inside a container with class containing "editor_project" or "editor_container"

        editor_clicked = False
        editor_selectors = [
            '[class*="editor_project"]',
            '[class*="editor_container"]',
            '.ace_editor',
            '.monaco-editor',
            '.CodeMirror',
        ]

        for selector in editor_selectors:
            try:
                editor = await page.query_selector(selector)
                if editor:
                    await editor.click()
                    await asyncio.sleep(0.3)
                    logger.info(f"[EBox] Clicked on editor: {selector}")
                    editor_clicked = True
                    break
            except:
                continue

        if not editor_clicked:
            logger.warning("[EBox] Could not click on editor container, trying direct keyboard input")

        # Now type the code using keyboard (this works with ACE/Monaco editors)
        # First select all existing code and delete it
        await page.keyboard.press("Control+a")
        await asyncio.sleep(0.1)
        await page.keyboard.press("Delete")
        await asyncio.sleep(0.1)

        # Type the code line by line (more reliable than pasting all at once)
        code_lines = code.split('\n')
        for i, line in enumerate(code_lines):
            await page.keyboard.type(line, delay=5)  # Small delay for each character
            if i < len(code_lines) - 1:
                await page.keyboard.press("Enter")
            # Log progress for long code
            if i % 10 == 0 and i > 0:
                logger.info(f"[EBox] Typed {i}/{len(code_lines)} lines...")

        logger.info(f"[EBox] Code entered ({len(code_lines)} lines)")
        await asyncio.sleep(0.5)

        # Fill command line arguments field if provided
        # Based on Playwright exploration: There IS a textbox labeled "Command line argument"
        # Located below the code editor - it loads dynamically after a delay
        if command_args:
            logger.info(f"[EBox] Filling command line arguments: {command_args}")
            try:
                # Wait for the Submit button to appear (indicates page is fully loaded)
                try:
                    await page.wait_for_selector('button:has-text("Submit")', timeout=5000)
                    logger.info("[EBox] Submit button found - page fully loaded")
                except:
                    logger.warning("[EBox] Submit button not found, continuing anyway")

                # Additional wait for dynamic content
                await asyncio.sleep(3)

                # Method 1: Use getByRole to find textbox (most reliable based on Playwright plugin test)
                try:
                    # In the Playwright plugin test, I used .first() which worked!
                    # Try the FIRST textbox instead of the last
                    textboxes = page.get_by_role('textbox')
                    count = await textboxes.count()
                    if count > 0:
                        # Try first textbox (the working one based on Playwright plugin test)
                        try:
                            await textboxes.first.fill(command_args)
                            logger.info(f"[EBox] ✅ Filled command args using first textbox: {command_args}")
                            return {"success": True}
                        except Exception as e1:
                            logger.warning(f"[EBox] First textbox failed: {e1}, trying last...")
                            # Fallback: try last textbox
                            try:
                                await textboxes.nth(count - 1).fill(command_args)
                                logger.info(f"[EBox] ✅ Filled command args using last textbox: {command_args}")
                                return {"success": True}
                            except Exception as e2:
                                logger.warning(f"[EBox] Last textbox also failed: {e2}")
                except Exception as e:
                    logger.warning(f"[EBox] getByRole method failed: {e}")

                # Method 2: Search through all inputs with broader selectors
                input_selectors = [
                    'input[type="text"]',
                    'input:not([type="password"]):not([type="hidden"])',
                    'input',
                    'textarea'
                ]

                for selector in input_selectors:
                    all_inputs = await page.query_selector_all(selector)
                    for inp in all_inputs:
                        try:
                            # Check parent, siblings, and grandparent for "command" + "argument"
                            parent_text = await inp.evaluate('el => el.parentElement ? el.parentElement.textContent : ""')
                            prev_sibling = await inp.evaluate('el => el.previousElementSibling ? el.previousElementSibling.textContent : ""')
                            grandparent = await inp.evaluate('el => el.parentElement && el.parentElement.parentElement ? el.parentElement.parentElement.textContent : ""')

                            combined = f"{parent_text} {prev_sibling} {grandparent}".lower()

                            if 'command' in combined and 'argument' in combined:
                                await inp.fill(command_args)
                                logger.info(f"[EBox] ✅ Filled command args: {command_args}")
                                return {"success": True}
                        except:
                            continue

                logger.warning("[EBox] Could not find command line argument input field")
            except Exception as e:
                logger.warning(f"[EBox] Error filling command args: {e}")

        await asyncio.sleep(1)

        # Click submit button (Using the exact class from deep analysis)
        logger.info("[EBox] Looking for Submit button (i-Design)...")
        submit_clicked = False
        submit_selectors = [
            'button.project_program_submit__36Qgn',  # The real E-Box submit button class
            'button.ui.button:has-text("Submit")',
            'button:has-text("Submit")',
            'button:has-text("Run")',
            'button:has-text("Execute")',
            'button:has-text("Check")',
        ]

        for selector in submit_selectors:
            try:
                btn = page.locator(selector).first
                if await btn.count() > 0:
                    await btn.click(force=True, timeout=3000)
                    logger.info(f"[EBox] Clicked submit button: {selector}")
                    submit_clicked = True
                    break
            except:
                continue

        if not submit_clicked:
            # Try JS fallback
            try:
                await page.evaluate("""() => {
                    const btns = Array.from(document.querySelectorAll('button'));
                    const sub = btns.find(b => b.textContent.trim().toLowerCase().includes('submit') || 
                                               b.className.includes('submit'));
                    if (sub) sub.click();
                }""")
                logger.info("[EBox] Clicked Submit via JS fallback")
                submit_clicked = True
            except:
                pass

        if not submit_clicked:
            return {"success": False, "error": "Could not find coding submit button"}

        # Wait for "Evaluation In Progress" modal to clear
        logger.info("[EBox] Waiting for evaluation to complete...")
        await asyncio.sleep(5)  # Wait for modal to appear and disappear
        
        # Check Previous Submission (a.item.tool_item__1NZ-N usually points to results)
        try:
            prev_sub = page.locator('a.item[class*="tool_item__"]').filter(has_text="Previous Submission").first
            if await prev_sub.count() > 0:
                logger.info("[EBox] Found 'Previous Submission' link - Evaluation likely finished.")
        except:
            pass

        return {"success": True, "message": "Code submitted and evaluation started"}

    except Exception as e:
        logger.error(f"[EBox] Submit error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


# ============================================================================
# i-LEARN: Reading Resource Handler
# ============================================================================

async def handle_learning_resource(browser) -> Dict[str, Any]:
    """
    Handle i-Learn 'Learning Resource' pages (PDF/document viewer).
    These are paginated documents — scroll through all pages then click Review.
    UI: Page numbers on left sidebar (1, 2, 3...), pink 'Review' button top-right.
    """
    try:
        page = browser.page
        logger.info("[EBox iLearn] Handling Learning Resource page...")
        await asyncio.sleep(2)

        pages_clicked = 0
        max_pages = 50  # safety limit

        for _ in range(max_pages):
            # Click the next page number in left sidebar (the first unread/unchecked one)
            next_page_result = await page.evaluate("""() => {
                // Find all page indicators in the left sidebar
                const pageItems = document.querySelectorAll('.slick-slide, .page-item, [class*="slide"], [class*="page_num"]');
                if (pageItems.length > 0) {
                    // Try to find and click the next unchecked page
                    for (const item of pageItems) {
                        const hasCheck = item.querySelector('[class*="check"], .green, [class*="complete"]');
                        if (!hasCheck) {
                            item.click();
                            return { clicked: true, text: item.innerText.trim() };
                        }
                    }
                }
                return { clicked: false };
            }""")

            await asyncio.sleep(1.5)
            pages_clicked += 1

            # Check if we reached the last page (Review button becomes prominent / enabled)
            review_visible = await page.evaluate("""() => {
                const btns = Array.from(document.querySelectorAll('button, a'));
                const reviewBtn = btns.find(b => b.textContent.trim().toLowerCase() === 'review');
                if (reviewBtn) {
                    const style = window.getComputedStyle(reviewBtn);
                    return {
                        found: true,
                        disabled: reviewBtn.disabled || style.opacity < 0.5,
                        text: reviewBtn.textContent.trim()
                    };
                }
                return { found: false };
            }""")

            if review_visible.get('found') and not review_visible.get('disabled'):
                logger.info(f"[EBox iLearn] Review button is active after {pages_clicked} pages")
                break

        # Click the pink Review button (top-right)
        review_clicked = False
        for sel in ['button:has-text("Review")', 'a:has-text("Review")', '.review-btn', 'button.review']:
            try:
                btn = page.locator(sel).first
                if await btn.count() > 0:
                    await btn.click(force=True, timeout=3000)
                    logger.info(f"[EBox iLearn] ✅ Clicked Review button via: {sel}")
                    review_clicked = True
                    break
            except Exception:
                continue

        if not review_clicked:
            # JS fallback
            try:
                await page.evaluate("""() => {
                    const btns = Array.from(document.querySelectorAll('button, a'));
                    const r = btns.find(b => b.textContent.trim().toLowerCase() === 'review');
                    if (r) r.click();
                }""")
                logger.info("[EBox iLearn] ✅ Clicked Review via JS fallback")
                review_clicked = True
            except Exception as e:
                logger.warning(f"[EBox iLearn] Review click failed: {e}")

        await asyncio.sleep(2)
        return {
            "success": True,
            "pages_read": pages_clicked,
            "review_clicked": review_clicked
        }

    except Exception as e:
        logger.error(f"[EBox iLearn] handle_learning_resource error: {e}")
        return {"success": False, "error": str(e)}


# ============================================================================
# i-DESIGN: File Upload Handler (AI-generated document)
# ============================================================================

async def generate_document_content(topic: str, task_description: str) -> str:
    """Use AWS Bedrock to generate a professional document for i-Design upload."""
    try:
        import boto3
        import json as _json

        bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')
        prompt = f"""You are a student completing a Communication Skills assignment.

TOPIC: {topic}
TASK: {task_description}

Write a well-structured, professional response document of 3-4 paragraphs.
The response should be relevant, grammatically correct, and demonstrate 
good communication skills.

Output ONLY the document text, no extra formatting or JSON."""

        response = bedrock.invoke_model(
            modelId='amazon.nova-pro-v1:0',
            body=_json.dumps({
                "messages": [{"role": "user", "content": [{"text": prompt}]}],
                "inferenceConfig": {"temperature": 0.4, "max_new_tokens": 800}
            })
        )
        body = _json.loads(response['body'].read())
        return body['output']['message']['content'][0]['text']
    except Exception as e:
        logger.error(f"[EBox iDesign] Doc generation error: {e}")
        # Fallback generic content
        return f"""Professional Communication Assignment – {topic}

Effective communication is a cornerstone of professional success. In the context of {topic}, 
it is essential to apply clear, concise, and purposeful language that conveys ideas accurately 
to the intended audience.

This assignment demonstrates an understanding of the key principles covered in this module. 
By reflecting on practical communication scenarios, one can develop greater competence in 
both written and verbal communication within professional settings.

In conclusion, mastering the principles of {topic} allows professionals to collaborate 
effectively, build trust, and contribute meaningfully to their organizations."""


async def handle_idesign_upload(browser, topic: str) -> Dict[str, Any]:
    """
    Handle i-Design file upload section in Communication Skills.
    Flow:
      1. Read the task description from the page
      2. Generate a .docx document using AWS Bedrock
      3. Upload using Playwright's set_input_files() on input#videoUpload
      4. Click Submit
    """
    try:
        import tempfile, os
        from docx import Document as DocxDocument

        page = browser.page
        logger.info(f"[EBox iDesign] Handling file upload for topic: {topic}")
        await asyncio.sleep(2)

        # ── Step 1: Read task description from page ───────────────────────
        body_text = await page.inner_text('body')
        # Try to find relevant task instructions (first 1000 chars is usually the brief)
        task_description = body_text[:1000].strip()
        logger.info(f"[EBox iDesign] Task snippet: {task_description[:100]}...")

        # ── Step 2: Generate document content using AI ────────────────────
        logger.info("[EBox iDesign] Generating document with AI...")
        doc_content = await generate_document_content(topic, task_description)

        # ── Step 3: Create .docx file ─────────────────────────────────────
        doc = DocxDocument()
        doc.add_heading(topic, 0)

        for para in doc_content.split('\n'):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        temp_path = os.path.join(tempfile.gettempdir(), f"ebox_idesign_{topic[:20].replace(' ', '_')}.docx")
        doc.save(temp_path)
        logger.info(f"[EBox iDesign] Saved document: {temp_path}")

        # ── Step 4: Upload via Playwright ─────────────────────────────────
        upload_success = False

        # Primary: E-Box uses input#videoUpload (discovered from DOM inspection)
        file_input_selectors = [
            'input#videoUpload',
            'input[type="file"]',
            'input[accept*="doc"]',
            'input[accept*="pdf"]',
        ]

        for sel in file_input_selectors:
            try:
                file_input = await page.query_selector(sel)
                if file_input:
                    await file_input.set_input_files(temp_path)
                    logger.info(f"[EBox iDesign] ✅ File set via: {sel}")
                    upload_success = True
                    await asyncio.sleep(2)
                    break
            except Exception as fe:
                logger.warning(f"[EBox iDesign] File input {sel} failed: {fe}")

        if not upload_success:
            # Try clicking "Select the file" button first to reveal the input
            try:
                select_btn = await page.query_selector('button.ui.black.basic.button')
                if select_btn:
                    await select_btn.click()
                    await asyncio.sleep(1)
                    # Now try the file input again
                    file_input = await page.query_selector('input[type="file"]')
                    if file_input:
                        await file_input.set_input_files(temp_path)
                        logger.info("[EBox iDesign] ✅ File set after clicking Select button")
                        upload_success = True
                        await asyncio.sleep(2)
            except Exception as be:
                logger.warning(f"[EBox iDesign] Select button fallback failed: {be}")

        if not upload_success:
            os.remove(temp_path)
            return {"success": False, "error": "Could not find file input element to upload"}

        # ── Step 5: Click Submit button ───────────────────────────────────
        submit_clicked = False
        for sel in [
            'button:has-text("Submit")',
            'button.ui.primary.button',
            'button[type="submit"]',
            '.submit-btn',
        ]:
            try:
                btn = page.locator(sel).first
                if await btn.count() > 0:
                    await btn.click(force=True, timeout=3000)
                    logger.info(f"[EBox iDesign] ✅ Clicked submit via: {sel}")
                    submit_clicked = True
                    break
            except Exception:
                continue

        if not submit_clicked:
            await page.evaluate("""() => {
                const btns = Array.from(document.querySelectorAll('button'));
                const s = btns.find(b => b.textContent.trim().toLowerCase().includes('submit'));
                if (s) s.click();
            }""")
            logger.info("[EBox iDesign] Clicked Submit via JS fallback")

        await asyncio.sleep(3)
        os.remove(temp_path)

        return {"success": True, "message": f"Document uploaded and submitted for topic: {topic}"}

    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        logger.error(f"[EBox iDesign] handle_idesign_upload error: {e}")
        return {"success": False, "error": str(e)}




async def extract_mcq_from_dom(page) -> Dict[str, Any]:
    """
    Extract the current MCQ question text and all answer options directly from the E-Box DOM.
    Uses the real class names discovered from live Playwright investigation:
      - Question container: div[class*='attempt_main_section__']
      - Answer options:    div[class*='mcq_answer__']  (role=radio)
    Returns dict with 'question', 'options' (list of texts), and 'raw_html'.
    """
    try:
        data = await page.evaluate("""() => {
            // ── Question text ───────────────────────────────────────────────
            const qContainer = document.querySelector(
                '[class*="attempt_main_section__"], [class*="question_container__"], ' +
                '[class*="question__"], .attempt_main_section'
            );
            let questionText = '';
            if (qContainer) {
                // grab first paragraph / heading that is NOT inside an answer div
                const qNode = qContainer.querySelector(
                    'p, h1, h2, h3, h4, [class*="question_text__"], [class*="qtext__"]'
                );
                if (qNode) questionText = qNode.innerText.trim();
                if (!questionText) {
                    // fallback: first non-empty text node of container (exclude option divs)
                    const answerDivs = qContainer.querySelectorAll('[class*="mcq_answer__"]');
                    answerDivs.forEach(a => a.remove());
                    questionText = qContainer.innerText.trim().split('\\n')[0];
                }
            }

            // ── Answer options ───────────────────────────────────────────────
            // Primary: div[role=radio] with class containing 'mcq_answer__'
            let optionEls = Array.from(
                document.querySelectorAll('[class*="mcq_answer__"][role="radio"]:not([class*="bar"]), [class*="mcq_answer__"]:not([class*="bar"])')
            );
            // Fallback: li or div with class containing 'option' or 'answer'
            if (optionEls.length === 0) {
                optionEls = Array.from(
                    document.querySelectorAll(
                        '[class*="option__"]:not([class*="bar"]):not([class*="nav"]):not([class*="sidebar"]):not([class*="menu"]), [class*="answer__"]:not([class*="bar"]):not([class*="nav"]):not([class*="sidebar"]):not([class*="menu"]), [class*="choice__"]'
                    )
                );
            }

            const options = optionEls.map((el, idx) => ({
                index: idx,
                text: el.innerText.trim(),
                classes: el.className,
                isGray: el.className.includes('gray'),
                isGreen: el.className.includes('green'),
                isRed:  el.className.includes('red'),
            }));

            // ── Question number from sidebar ─────────────────────────────────
            const sidebar = document.querySelector(
                '[class*="problem_list__"], [class*="question_nav__"]'
            );
            const activeQ = sidebar ? sidebar.querySelector('[class*="active__"]') : null;
            const qNum = activeQ ? activeQ.innerText.trim() : '?';

            return { questionText, options, questionNumber: qNum, optionCount: options.length };
        }""")
        return {"success": True, **data}
    except Exception as e:
        logger.error(f"[EBox MCQ] DOM extraction error: {e}")
        return {"success": False, "error": str(e), "options": [], "questionText": ""}


async def reason_with_bedrock(question_text: str, options: List[Dict]) -> Dict[str, Any]:
    """
    Use AWS Bedrock Nova Pro with full chain-of-thought reasoning to pick the correct MCQ answer.
    Returns { answer_index, answer_text, confidence, reasoning }.
    """
    import boto3, json as _json

    option_lines = "\n".join(
        f"  Option {i+1}: {opt['text']}" for i, opt in enumerate(options)
    )

    system_prompt = (
        "You are an expert academic tutor helping a student answer multiple-choice questions "
        "across Engineering Mathematics, Digital Electronics, Biology, Chemistry, Physics, "
        "Python programming, and C programming. "
        "Think step by step before giving your final answer."
    )

    user_prompt = f"""Answer this MCQ question correctly.

QUESTION:
{question_text}

OPTIONS:
{option_lines}

Instructions:
1. Read the question carefully.
2. Use your deep knowledge to evaluate EACH option.
3. Eliminate wrong answers with reasoning.
4. State which option is CORRECT and why.
5. Return ONLY valid JSON (no markdown fences):

{{
  "answer_index": <0-based integer index of correct option>,
  "answer_text": "<exact text of the correct option>",
  "confidence": "<high|medium|low>",
  "reasoning": "<concise chain-of-thought explanation>"
}}"""

    try:
        bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')
        response = bedrock.invoke_model(
            modelId='amazon.nova-pro-v1:0',
            body=_json.dumps({
                "system": [{"text": system_prompt}],
                "messages": [{"role": "user", "content": [{"text": user_prompt}]}],
                "inferenceConfig": {"temperature": 0.1, "max_new_tokens": 600}
            })
        )
        body = _json.loads(response['body'].read())
        raw = body['output']['message']['content'][0]['text']

        result = _extract_json_from_text(raw)
        logger.info(
            f"[EBox MCQ] Bedrock chose option {result.get('answer_index', '?')+1} "
            f"({result.get('confidence','?')} confidence): {result.get('reasoning','')[:120]}"
        )
        return {"success": True, **result}

    except Exception as e:
        logger.error(f"[EBox MCQ] Bedrock reasoning failed: {e}")
        # Fallback: pick option 0 and warn
        return {
            "success": False,
            "answer_index": 0,
            "answer_text": options[0]['text'] if options else "",
            "confidence": "low",
            "reasoning": f"Bedrock error: {e}",
        }


async def click_mcq_option(page, answer_index: int, answer_text: str) -> Dict[str, Any]:
    """
    Click the correct MCQ answer option on E-Box using the real DOM structure.
    Primary: click the Nth div[class*='mcq_answer__'] by index.
    Fallback: match by text content.
    Returns { clicked: bool, method: str }.
    """
    # ── Method 1: Click by index on real mcq_answer divs ─────────────────────
    try:
        clicked = await page.evaluate(f"""(targetIdx) => {{
            const options = Array.from(
                document.querySelectorAll('[class*="mcq_answer__"][role="radio"]:not([class*="bar"]), [class*="mcq_answer__"]:not([class*="bar"])')
            );
            if (options.length === 0) return {{ ok: false, reason: 'no mcq_answer divs found' }};
            const idx = Math.min(targetIdx, options.length - 1);
            options[idx].click();
            return {{ ok: true, method: 'index-click on mcq_answer div', count: options.length, clicked: idx }};
        }}""", answer_index)
        if clicked.get('ok'):
            logger.info(f"[EBox MCQ] ✅ Clicked option {answer_index+1} via mcq_answer div index")
            await asyncio.sleep(0.5)
            return {"clicked": True, "method": clicked.get('method')}
    except Exception as e:
        logger.warning(f"[EBox MCQ] Method 1 (index click) failed: {e}")

    # ── Method 2: Match by option text ───────────────────────────────────────
    if answer_text:
        try:
            clicked = await page.evaluate(f"""(targetText) => {{
                const options = Array.from(
                    document.querySelectorAll('[class*="mcq_answer__"], [class*="option__"]:not([class*="bar"]):not([class*="nav"]):not([class*="sidebar"]):not([class*="menu"]), [class*="choice__"]')
                );
                const norm = s => s.trim().toLowerCase().replace(/\\s+/g, ' ');
                const tgt = norm(targetText);
                // exact match first
                let el = options.find(o => norm(o.innerText) === tgt);
                // partial match fallback
                if (!el) el = options.find(o => norm(o.innerText).includes(tgt) || tgt.includes(norm(o.innerText)));
                if (el) {{ el.click(); return {{ ok: true, method: 'text-match click' }}; }}
                return {{ ok: false, reason: 'text not found among ' + options.length + ' options' }};
            }}""", answer_text)
            if clicked.get('ok'):
                logger.info(f"[EBox MCQ] ✅ Clicked option by text match: '{answer_text[:60]}'")
                await asyncio.sleep(0.5)
                return {"clicked": True, "method": "text-match"}
        except Exception as e:
            logger.warning(f"[EBox MCQ] Method 2 (text match) failed: {e}")

    # ── Method 3: Playwright locator force-click ──────────────────────────────
    try:
        opts = page.locator('[class*="mcq_answer__"]:not([class*="bar"])')
        count = await opts.count()
        if count > 0:
            idx = min(answer_index, count - 1)
            await opts.nth(idx).click(force=True, timeout=3000)
            logger.info(f"[EBox MCQ] ✅ Clicked option {idx+1} via Playwright force-click")
            await asyncio.sleep(0.5)
            return {"clicked": True, "method": "playwright-force"}
    except Exception as e:
        logger.warning(f"[EBox MCQ] Method 3 (Playwright locator) failed: {e}")

    logger.error("[EBox MCQ] ❌ All click methods failed")
    return {"clicked": False, "method": "none"}


async def verify_selection_and_review(page) -> Dict[str, Any]:
    """
    After clicking an answer, click 'Review Answer' and verify whether the selection
    turned green (correct) or red (wrong) using E-Box's real CSS class names.
    Returns { clicked_review: bool, is_correct: bool, correct_option_text: str }.
    """
    # ── Click Review Answer ────────────────────────────────────────────────
    review_clicked = False
    review_selectors = [
        'button.ui.button:not(.primary):not(.secondary)',  # "Review Answer" is plain ui.button
        'button:has-text("Review Answer")',
        'button:has-text("Review")',
    ]
    for sel in review_selectors:
        try:
            btn = page.locator(sel).first
            if await btn.count() > 0:
                await btn.click(force=True, timeout=3000)
                logger.info(f"[EBox MCQ] Clicked Review Answer via: {sel}")
                review_clicked = True
                break
        except Exception:
            continue

    if not review_clicked:
        # JS fallback
        try:
            await page.evaluate("""() => {
                const btns = Array.from(document.querySelectorAll('button'));
                const rb = btns.find(b => b.textContent.trim().toLowerCase().includes('review'));
                if (rb) rb.click();
            }""")
            review_clicked = True
            logger.info("[EBox MCQ] Clicked Review Answer via JS fallback")
        except Exception as e:
            logger.warning(f"[EBox MCQ] Review Answer click failed entirely: {e}")

    await asyncio.sleep(2)  # Wait for color feedback to appear

    # ── Check color state of selected option ─────────────────────────────
    state = await page.evaluate("""() => {
        const options = Array.from(document.querySelectorAll('[class*="mcq_answer__"]'));
        const selected = options.find(o => o.className.includes('green') || o.className.includes('red'));
        if (!selected) return { found: false };
        const isCorrect = selected.className.includes('green');
        // Also find the green (correct) answer if we got it wrong
        const correctEl = options.find(o => o.className.includes('green'));
        return {
            found: true,
            isCorrect,
            selectedClass: selected.className,
            correctText: correctEl ? correctEl.innerText.trim() : '',
        };
    }""")

    is_correct = state.get('isCorrect', False)
    correct_text = state.get('correctText', '')
    logger.info(
        f"[EBox MCQ] Review result: {'✅ CORRECT' if is_correct else '❌ WRONG'} "
        f"| correct answer: '{correct_text[:60]}'"
    )

    return {
        "clicked_review": review_clicked,
        "is_correct": is_correct,
        "correct_option_text": correct_text,
        "state_found": state.get('found', False),
    }


async def click_next_question(page) -> Dict[str, Any]:
    """
    Click the Next button to proceed to the next MCQ question.
    E-Box uses button.ui.primary.button for Next.
    Returns { clicked: bool, is_last: bool }.
    """
    next_selectors = [
        'button.ui.primary.button',           # real E-Box Next button class
        'button:has-text("Next")',
        'button:has-text("Save And Close")',
        'button:has-text("Save & Close")',
        'button:has-text("Finish")',
    ]
    for sel in next_selectors:
        try:
            btn = page.locator(sel).first
            if await btn.count() > 0:
                btn_text = (await btn.inner_text()).strip().lower()
                await btn.click(force=True, timeout=3000)
                logger.info(f"[EBox MCQ] Clicked '{btn_text}' button")
                await asyncio.sleep(1.5)
                is_last = any(k in btn_text for k in ['save', 'close', 'finish'])
                return {"clicked": True, "is_last": is_last}
        except Exception:
            continue

    # JS fallback
    try:
        is_last = await page.evaluate("""() => {
            const btns = Array.from(document.querySelectorAll('button'));
            const nb = btns.find(b => /next|save|finish/i.test(b.textContent));
            if (nb) { nb.click(); return /save|finish/i.test(nb.textContent); }
            return null;
        }""")
        if is_last is not None:
            logger.info("[EBox MCQ] Clicked Next/Save via JS fallback")
            await asyncio.sleep(1.5)
            return {"clicked": True, "is_last": bool(is_last)}
    except Exception as e:
        logger.warning(f"[EBox MCQ] Next button JS fallback failed: {e}")

    return {"clicked": False, "is_last": False}


async def solve_mcq_question(browser) -> Dict[str, Any]:
    """
    Solve a single MCQ question on the current E-Box attempt page.
    Pipeline:
      1. Extract question + options from real DOM (no vision required for reading)
      2. Send to AWS Bedrock Nova Pro for chain-of-thought reasoning
      3. Click the selected answer using real mcq_answer__ CSS class selectors
      4. Click Review Answer and verify green/red state
    Returns full result dict.
    """
    try:
        page = browser.page
        await asyncio.sleep(1)

        # ── Step 1: Extract question from DOM ─────────────────────────────
        dom_data = await extract_mcq_from_dom(page)
        question_text = dom_data.get("questionText", "").strip()
        options = dom_data.get("options", [])
        q_num = dom_data.get("questionNumber", "?")

        logger.info(f"[EBox MCQ] Q#{q_num}: {question_text[:100]}…")
        logger.info(f"[EBox MCQ] Options found: {len(options)}")
        for opt in options:
            logger.info(f"  [{opt['index']+1}] {opt['text'][:80]}")

        if not options:
            logger.warning("[EBox MCQ] No options found in DOM — taking screenshot for vision fallback")
            # Vision fallback: take screenshot and use Bedrock multimodal
            import base64, boto3, json as _json
            screenshot_b64 = base64.b64encode(await page.screenshot()).decode()
            bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')
            resp = bedrock.invoke_model(
                modelId='amazon.nova-pro-v1:0',
                body=_json.dumps({
                    "messages": [{"role": "user", "content": [
                        {"image": {"format": "png", "source": {"bytes": screenshot_b64}}},
                        {"text": (
                            "This is an E-Box MCQ question screenshot. "
                            "Identify the question and all answer options. "
                            "Return JSON: {\"question\": \"...\", "
                            "\"options\": [\"opt1\",\"opt2\",...], "
                            "\"answer_index\": 0, \"answer_text\": \"...\", \"reasoning\": \"...\"}"
                        )}
                    ]}],
                    "inferenceConfig": {"temperature": 0.1, "max_new_tokens": 600}
                })
            )
            raw = _json.loads(resp['body'].read())['output']['message']['content'][0]['text']
            vision_result = _extract_json_from_text(raw)
            answer_index = vision_result.get("answer_index", 0)
            answer_text  = vision_result.get("answer_text", "")
            reasoning    = vision_result.get("reasoning", "")
            options = [{"index": i, "text": t} for i, t in enumerate(vision_result.get("options", []))]
        else:
            # ── Step 2: Bedrock reasoning ─────────────────────────────────
            ai = await reason_with_bedrock(question_text, options)
            answer_index = ai.get("answer_index", 0)
            answer_text  = ai.get("answer_text", options[answer_index]["text"] if options else "")
            reasoning    = ai.get("reasoning", "")

        # ── Step 3: Click the answer ──────────────────────────────────────
        click_result = await click_mcq_option(page, answer_index, answer_text)

        return {
            "success": click_result["clicked"],
            "answer_index": answer_index,
            "answer_text": answer_text,
            "question_text": question_text,
            "reasoning": reasoning,
            "click_method": click_result.get("method", "none"),
        }

    except Exception as e:
        logger.error(f"[EBox MCQ] solve_mcq_question error: {e}")
        import traceback; traceback.print_exc()
        return {"success": False, "error": str(e)}




async def complete_mcq_section(browser, topic_name: str, section_name: str) -> Dict[str, Any]:
    """
    Complete ALL MCQ questions in i-Learn, i-Explore, or i-Analyse sections.
    Flow per question:
      1. solve_mcq_question  → extract DOM + Bedrock reasoning + click answer
      2. verify_selection_and_review → click Review Answer, check green/red
      3. click_next_question → move to next or detect end
    Uses real E-Box DOM selectors (mcq_answer__, ui.primary.button, etc.)
    """
    results = []
    questions_completed = 0
    correct_count = 0
    max_questions = 60  # Safety limit

    try:
        page = browser.page
        logger.info(f"[EBox MCQ] ══ Starting section: {topic_name} / {section_name} ══")

        # ── Navigate to the section tab ──────────────────────────────────────
        section_result = await click_section(browser, section_name)
        if not section_result.get("success"):
            return {"success": False, "error": f"Could not click section tab '{section_name}'"}
        await asyncio.sleep(2)

        # ── Find and click the attempt link ──────────────────────────────────
        attempt_links = await page.query_selector_all('a[href*="attempt"]')
        if not attempt_links:
            return {"success": False, "error": "No attempt link found in this section"}
        await attempt_links[0].click()
        await asyncio.sleep(3)
        logger.info(f"[EBox MCQ] Entered attempt page: {await browser.get_current_url()}")

        # ── Detect page type ──────────────────────────────────────────────────
        page_type = await page.evaluate("""() => {
            const body = document.body.innerText.toLowerCase();
            // Learning Resource: has paginated PDF viewer
            const hasLearningResource = document.querySelector('[class*="slide"], [class*="page_num"], .slick-slider') !== null;
            // MCQ: has mcq_answer divs
            const hasMcq = document.querySelector('[class*="mcq_answer__"]') !== null;
            // File upload: has file input
            const hasFileUpload = document.querySelector('input[type="file"], input#videoUpload') !== null;
            // Also check text indicators
            const isLearningText = body.includes('learning resource') || body.includes('module ');
            return { hasLearningResource, hasMcq, hasFileUpload, isLearningText };
        }""")

        logger.info(f"[EBox] Page type detection: {page_type}")

        # Route: Learning Resource (PDF reader)
        if page_type.get('hasLearningResource') or page_type.get('isLearningText'):
            if not page_type.get('hasMcq') and not page_type.get('hasFileUpload'):
                logger.info("[EBox] ➡ Detected Learning Resource — calling handle_learning_resource")
                result = await handle_learning_resource(browser)
                return {
                    "success": result.get("success", False),
                    "problems_completed": 1 if result.get("success") else 0,
                    "results": [f"📖 Reading resource completed: {result.get('pages_read', 0)} pages"],
                    "type": "learning_resource"
                }

        # Route: File Upload (i-Design)
        if page_type.get('hasFileUpload'):
            logger.info("[EBox] ➡ Detected File Upload — calling handle_idesign_upload")
            result = await handle_idesign_upload(browser, topic_name)
            return {
                "success": result.get("success", False),
                "problems_completed": 1 if result.get("success") else 0,
                "results": [f"📎 File uploaded: {result.get('message', result.get('error', ''))}"],
                "type": "file_upload"
            }

        # Default: MCQ loop ─────────────────────────────────────────────────────

        for q_num in range(1, max_questions + 1):
            logger.info(f"[EBox MCQ] ── Question #{q_num} ──")
            await asyncio.sleep(1)

            # Check if we're still on an attempt page
            current_url = await browser.get_current_url()
            if "attempt" not in current_url:
                logger.info("[EBox MCQ] Left attempt page — section complete")
                break

            # Step 1: Solve (extract + reason + click)
            solve = await solve_mcq_question(browser)
            if not solve.get("success"):
                results.append(f"❌ Q{q_num}: solve failed — {solve.get('error','')}")
                # Click Review Answer to save whatever was selected
                await verify_selection_and_review(page)
                next_r = await click_next_question(page)
                if next_r.get("is_last") or not next_r.get("clicked"):
                    # No Next — save via Review Answer and stop
                    await verify_selection_and_review(page)
                    logger.info("[EBox MCQ] No Next after failed solve — saved via Review Answer")
                    break
                continue

            # Step 2: Review Answer + verify color
            review = await verify_selection_and_review(page)
            is_correct = review.get("is_correct", False)
            correct_text = review.get("correct_option_text", "")
            if is_correct:
                correct_count += 1
                icon = "✅"
            else:
                icon = "❌"

            result_line = (
                f"{icon} Q{q_num}: Answered '{solve['answer_text'][:50]}' "
                f"| {'Correct' if is_correct else f'Wrong — right: {correct_text[:40]}'}"
            )
            logger.info(f"[EBox MCQ] {result_line}")
            results.append(result_line)
            questions_completed += 1

            # Step 3: Try Next question
            await asyncio.sleep(0.5)
            next_r = await click_next_question(page)
            if not next_r.get("clicked"):
                # Next button absent — click Review Answer to save progress
                logger.info("[EBox MCQ] Next not found — clicking Review Answer to save and ending section")
                await verify_selection_and_review(page)
                break
            if next_r.get("is_last"):
                logger.info("[EBox MCQ] Last question — section complete")
                break

            await asyncio.sleep(1.5)

        logger.info(
            f"[EBox MCQ] Section complete: {questions_completed} answered, "
            f"{correct_count}/{questions_completed} correct"
        )
        return {
            "success": True,
            "problems_completed": questions_completed,
            "correct": correct_count,
            "results": results,
        }

    except Exception as e:
        logger.error(f"[EBox MCQ] complete_mcq_section error: {e}")
        import traceback; traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "problems_completed": questions_completed,
            "results": results,
        }



async def complete_section(browser, topic_name: str, section_name: str) -> Dict[str, Any]:
    """
    Complete ALL problems/projects within a specific section.
    Returns results for all problems attempted.
    """
    results = []
    problems_completed = 0
    max_problems = 20  # Safety limit to prevent infinite loops

    try:
        page = browser.page
        logger.info(f"[EBox] Starting section: {topic_name}/{section_name}")

        # Click on section tab
        section_result = await click_section(browser, section_name)
        if not section_result.get("success"):
            return {"success": False, "error": f"Could not click section {section_name}"}

        await asyncio.sleep(2)

        # Loop through all problems in this section
        for problem_num in range(1, max_problems + 1):
            logger.info(f"[EBox] Problem #{problem_num} in {section_name}")

            # Get all project/attempt links on the current page
            attempt_links = await page.query_selector_all('a[href*="attempt"]')

            if not attempt_links or len(attempt_links) == 0:
                logger.info(f"[EBox] No more problems found in {section_name}")
                break

            # Click the first unfinished problem
            first_link = attempt_links[0]
            link_text = await first_link.inner_text()
            href = await first_link.get_attribute('href')

            logger.info(f"[EBox] Clicking problem: {link_text[:50]}... ({href})")
            await first_link.click()
            await asyncio.sleep(3)

            # Analyze the problem
            problem_result = await analyze_problem_page(browser)
            if not problem_result.get("success"):
                results.append(f"❌ Problem {problem_num}: Could not analyze")
                # Go back and try next problem
                await page.go_back()
                await asyncio.sleep(2)
                continue

            problem = problem_result["problem"]

            # Generate solution
            solution_result = await generate_solution_code(problem)
            if not solution_result.get("success"):
                results.append(f"❌ Problem {problem_num}: Could not generate solution")
                await page.go_back()
                await asyncio.sleep(2)
                continue

            # Submit code
            submit_result = await submit_code(
                browser,
                solution_result["code"],
                solution_result.get("command_args", "")
            )

            if submit_result.get("success"):
                results.append(f"✅ Problem {problem_num}: {problem.get('problem_title', 'Unknown')[:40]}...")
                problems_completed += 1
            else:
                results.append(f"⚠️ Problem {problem_num}: Submit failed")

            # Wait for submission to process
            await asyncio.sleep(3)

            # Go back to section page to get next problem
            await page.go_back()
            await asyncio.sleep(2)

            # Check if we're still on the section page
            current_url = await browser.get_current_url()
            if 'attempt' in current_url:
                # Still on a problem page, go back again
                await page.go_back()
                await asyncio.sleep(2)

        logger.info(f"[EBox] Section {section_name} complete: {problems_completed} problems solved")

        return {
            "success": True,
            "problems_completed": problems_completed,
            "results": results
        }

    except Exception as e:
        logger.error(f"[EBox] Error completing section: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "problems_completed": problems_completed,
            "results": results
        }


@function_tool()
async def complete_ebox_course(
    context: RunContext,
    request: Annotated[str, "Natural language request like 'finish differential equations', 'complete biology course', 'do i-Design section'"]
) -> str:
    """
    Automatically complete E-Box course projects/problems.

    E-Box Structure:
    - Course → Topics (left sidebar) → Sections (i-Learn, i-Explore, i-Analyse, i-Design) → Projects → Problems

    Supports:
    - "finish differential equations" - complete all topics and sections
    - "do i-Design section" - complete ALL problems in i-Design across all topics
    - "solve vector calculus" - complete all sections in vector calculus topic
    - "finish i-Learn" - complete i-Learn section in all topics
    """
    browser = None
    try:
        from browser_automation import BrowserAutomationEngine

        parsed = extract_course_and_topic(request)
        logger.info(f"[EBox] Parsed request: {parsed}")

        # Start browser
        browser = BrowserAutomationEngine(backend="playwright", headless=False)
        await browser.start()
        logger.info("[EBox] Browser started")

        # Login
        login_result = await login_to_ebox(browser)
        if not login_result.get("success"):
            return f"Login failed: {login_result.get('error')}. Browser left open for manual login."

        # Navigate to course
        if parsed["course"]:
            course_result = await navigate_to_course(browser, parsed["course"])
            if not course_result.get("success"):
                return f"Could not find course: {parsed['course']}. Browser left open."

        await asyncio.sleep(2)

        # Get list of topics
        topics_result = await get_topics_list(browser)
        topics = topics_result.get("topics", [])

        if not topics:
            return "No topics found on the course page. Please navigate manually. Browser left open."

        all_results = []
        total_problems = 0

        # Process each topic
        for topic in topics:
            logger.info(f"[EBox] ========== TOPIC: {topic} ==========")

            # Click on the topic
            topic_result = await click_topic(browser, topic)
            if not topic_result.get("success"):
                all_results.append(f"❌ {topic}: Could not select topic")
                continue

            # Process each section
            sections_to_do = [parsed["section"]] if parsed["section"] else SECTIONS

            for section in sections_to_do:
                logger.info(f"[EBox] ----- SECTION: {section} -----")

                # Route based on section type
                # For Technical Communication, all sections are MCQs or Reading
                is_mcq_course = parsed["course"] == "Technical Communication"
                mcq_sections = ["i-learn", "ilearn", "i-explore", "iexplore", "i-assess", "iassess"]
                
                if is_mcq_course or section.lower() in mcq_sections:
                    logger.info(f"[EBox] Detected MCQ/Quiz/Reading section: {section}")
                    section_result = await complete_mcq_section(browser, topic, section)
                else:
                    logger.info(f"[EBox] Detected Coding/Project section: {section}")
                    section_result = await complete_section(browser, topic, section)

                if section_result.get("success"):
                    problems_count = section_result.get("problems_completed", 0)
                    total_problems += problems_count
                    all_results.append(f"✅ {topic}/{section}: {problems_count} problems completed")

                    # Add individual problem results
                    for result in section_result.get("results", []):
                        all_results.append(f"  {result}")
                else:
                    all_results.append(f"❌ {topic}/{section}: {section_result.get('error', 'Unknown error')}")

        # Build response
        response = f"**E-Box Course Automation Complete!**\n\n"
        response += f"**Total Problems Solved: {total_problems}**\n\n"
        response += f"**Results:**\n"
        for r in all_results:
            response += f"{r}\n"

        response += f"\n**Browser left open for verification.**"

        return response

    except Exception as e:
        logger.error(f"[EBox] Error: {e}")
        import traceback
        traceback.print_exc()
        return f"Automation error: {str(e)}. Browser left open for manual recovery."


@function_tool()
async def list_ebox_courses(context: RunContext) -> str:
    """
    List all available E-Box courses and their topics for the current user (SIT25CS170).
    Use this when the user asks 'what courses do I have on e-box', 'show my courses', etc.
    No browser needed — returns the pre-loaded catalogue instantly.
    """
    lines = ["**Your E-Box Courses:**\n"]
    for i, (course, topics) in enumerate(COURSE_TOPICS.items(), 1):
        lines.append(f"**{i}. {course}**")
        for j, topic in enumerate(topics, 1):
            lines.append(f"   {j}. {topic}")
        lines.append("")
    lines.append(f"**Sections in each topic:** {', '.join(SECTIONS)}")
    lines.append("\nSay 'complete [course name]' or 'finish [topic]' to start automation.")
    return "\n".join(lines)


@function_tool()
async def ebox_open_course(
    context: RunContext,
    course: Annotated[str, "Course name or alias, e.g. 'differential', 'python', 'digital electronics'"],
) -> str:
    """
    Open E-Box, log in, and navigate directly to a specific course page.
    Use this when the user wants to open or view a course without full automation.
    """
    browser = None
    try:
        from browser_automation import BrowserAutomationEngine

        # Resolve alias
        resolved = COURSE_MAPPINGS.get(course.strip().lower(), course)
        logger.info(f"[EBox] Opening course: {resolved}")

        browser = BrowserAutomationEngine(backend="playwright", headless=False)
        await browser.start()

        login_result = await login_to_ebox(browser)
        if not login_result.get("success"):
            return f"❌ Login failed: {login_result.get('error')}"

        course_result = await navigate_to_course(browser, resolved)
        if not course_result.get("success"):
            return (
                f"✅ Logged in successfully!\n"
                f"❌ Could not navigate to '{resolved}' automatically — dashboard is open in the browser.\n"
                f"Available courses: {', '.join(COURSE_TOPICS.keys())}"
            )

        # Get topics visible on the page
        topics_result = await get_topics_list(browser)
        topics = topics_result.get("topics", [])

        url = course_result.get("url", "")
        topic_str = "\n".join(f"  • {t}" for t in topics) if topics else "  (could not read topics from page)"

        return (
            f"✅ Opened **{resolved}** in the browser!\n"
            f"URL: {url}\n\n"
            f"**Topics found:**\n{topic_str}\n\n"
            f"Say 'complete this course' or 'do i-Design section' to begin automation."
        )

    except Exception as e:
        logger.error(f"[EBox] ebox_open_course error: {e}")
        return f"Error: {e}"


@function_tool()
async def ebox_help_with_problem(
    context: RunContext,
    problem_description: Annotated[str, "Describe the problem or paste the problem statement"]
) -> str:
    """
    Get AI help with understanding and solving an E-Box problem.
    Returns solution approach, code, and explanation.
    """
    try:
        import boto3
        import json

        bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')

        prompt = f"""
You are an expert programmer helping with an E-Box course problem.

PROBLEM:
{problem_description}

Provide:
1. Understanding of what the problem is asking
2. Solution approach (step-by-step)
3. Complete code solution (if applicable)
4. Explanation of key concepts

Return JSON:
{{
    "understanding": "what the problem is asking for",
    "approach": ["step 1", "step 2", ...],
    "code": "complete code solution (if applicable)",
    "key_concepts": ["concept 1", "concept 2", ...],
    "tips": "helpful tips for solving similar problems"
}}

Return ONLY valid JSON.
"""

        response = bedrock.invoke_model(
            modelId='amazon.nova-pro-v1:0',
            body=json.dumps({
                "messages": [{"role": "user", "content": prompt}],
                "inferenceConfig": {
                    "temperature": 0.4,
                    "max_new_tokens": 2000
                }
            })
        )

        response_body = json.loads(response['body'].read())
        result_text = response_body['output']['message']['content'][0]['text']

        result = _extract_json_from_text(result_text)

        # Format response
        output = f"**Understanding:** {result.get('understanding', '')}\n\n"
        output += f"**Solution Approach:**\n"
        for i, step in enumerate(result.get('approach', []), 1):
            output += f"{i}. {step}\n"

        if result.get('code'):
            output += f"\n**Code Solution:**\n```\n{result['code']}\n```\n"

        output += f"\n**Key Concepts:** {', '.join(result.get('key_concepts', []))}\n"
        output += f"\n**Tips:** {result.get('tips', '')}"

        return output

    except Exception as e:
        logger.error(f"[EBox] Help error: {e}")
        return f"Could not generate help: {str(e)}"
